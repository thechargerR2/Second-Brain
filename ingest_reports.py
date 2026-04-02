"""
Ingest co-created reports from stock_screener/reports/ into Second Brain.
Handles .txt, .html, .docx, .xlsx, .pdf files.
Skips duplicates (checks by title), skips temp files and recurring reports.
"""

import os
import re
import sqlite3
from pathlib import Path

REPORTS_DIR = "/Users/ronrusso/stock_screener/reports"
DB_PATH = "/Users/ronrusso/second-brain/second_brain.db"

# Skip recurring/auto-generated reports and temp files
SKIP_PATTERNS = [
    r"^screener_",
    r"^hotlist_",
    r"^daily_report_",
    r"^morning_briefing_",
    r"^portfolio_performance_",
    r"^dd_memo_[A-Z]",       # the generator script, not the output
    r"^~\$",                  # Word temp files
    r"_latest\.",             # symlinks
]


def should_skip(filename):
    for pat in SKIP_PATTERNS:
        if re.match(pat, filename, re.IGNORECASE):
            return True
    return False


def classify(filename):
    """Guess a topic from the filename."""
    fn = filename.lower()
    if "cinqcare" in fn or "access" in fn or "cinqiq" in fn or "mentorship" in fn:
        return "CINQCARE / Healthcare"
    if "dd_" in fn or "due_diligence" in fn or "broadburch_partners_dd" in fn:
        return "Due Diligence"
    if "robot" in fn or "figure_ai" in fn:
        return "Robotics & AI"
    if "space" in fn or "satellite" in fn:
        return "Space & Infrastructure"
    if "portfolio" in fn or "rebalanc" in fn or "dividend" in fn or "investment_memo" in fn or "growth_stock" in fn:
        return "Portfolio & Strategy"
    if "claremont" in fn or "detention" in fn or "185_west_gale" in fn or "govt_programs" in fn:
        return "Real Estate & Government"
    if "dc_picks" in fn or "private_credit" in fn or "rwr_" in fn:
        return "Investment Research"
    if "brad_smith" in fn or "function_health" in fn:
        return "Investment Research"
    if "broadburch_capital" in fn:
        return "Portfolio & Strategy"
    return "Research"


def read_txt(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def read_html(path):
    from bs4 import BeautifulSoup
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    return soup.get_text(separator="\n", strip=True)


def read_docx(path):
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also grab tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def read_pdf(path):
    from PyPDF2 import PdfReader
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n".join(parts)


def read_file(path):
    ext = path.suffix.lower()
    try:
        if ext == ".txt":
            return read_txt(path)
        elif ext == ".html":
            return read_html(path)
        elif ext == ".docx":
            return read_docx(path)
        elif ext == ".xlsx":
            return read_xlsx(path)
        elif ext == ".pdf":
            return read_pdf(path)
        else:
            return None
    except Exception as e:
        print(f"  ERROR reading {path.name}: {e}")
        return None


def make_title(filename):
    """Convert filename to a readable title."""
    name = Path(filename).stem
    # Remove date suffixes like _2026-02-15 or _20260215 or _2026-02-15_1841
    name = re.sub(r"_\d{4}-\d{2}-\d{2}(_\d{4})?$", "", name)
    name = re.sub(r"_\d{8}$", "", name)
    # Replace underscores with spaces
    name = name.replace("_", " ")
    return name.strip()


def main():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row

    # Get existing titles to avoid duplicates
    existing = set()
    for row in conn.execute("SELECT title FROM entries"):
        existing.add(row["title"].lower().strip())

    reports_path = Path(REPORTS_DIR)
    files = sorted(reports_path.iterdir())

    ingested = 0
    skipped_dup = 0
    skipped_filter = 0
    errors = 0

    for f in files:
        if not f.is_file():
            continue
        if f.suffix.lower() not in (".txt", ".html", ".docx", ".xlsx", ".pdf"):
            continue
        if should_skip(f.name):
            skipped_filter += 1
            continue
        if f.name.endswith("_latest.docx") or f.name.endswith("_latest.xlsx"):
            skipped_filter += 1
            continue

        title = make_title(f.name)

        # Check for duplicates (by title, case-insensitive)
        if title.lower() in existing:
            skipped_dup += 1
            continue

        content = read_file(f)
        if content is None:
            errors += 1
            continue

        # Truncate very large content (keep first 50K chars)
        if len(content) > 50000:
            content = content[:50000] + "\n\n[... truncated ...]"

        topic = classify(f.name)

        conn.execute(
            "INSERT INTO entries (type, title, content, topic, url) VALUES (?, ?, ?, ?, ?)",
            ("document", title, content, topic, ""),
        )
        existing.add(title.lower())
        ingested += 1
        print(f"  + {title} [{topic}]")

    conn.commit()
    conn.close()

    print(f"\nDone! Ingested: {ingested}, Skipped (duplicate): {skipped_dup}, "
          f"Skipped (filter): {skipped_filter}, Errors: {errors}")


if __name__ == "__main__":
    main()
